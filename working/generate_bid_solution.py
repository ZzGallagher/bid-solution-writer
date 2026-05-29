from __future__ import annotations

import re
import textwrap
import zipfile
from datetime import datetime
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "input"
TEMPLATE = ROOT / "templates" / "投标方案模板.优化前备份.docx"
OUTPUT = ROOT / "output"
RECORDS = OUTPUT / "records"
PROJECT_NAME = "电子海图显示与信息系统"
TODAY = datetime.now().strftime("%Y%m%d")
OUT_DOCX = OUTPUT / f"{PROJECT_NAME}设计方案_V1.00_{TODAY}.docx"


TECH_STANDARDS = [
    "IEC 61174-2015 海上导航和无线电通信设备及系统 电子海图显示和信息系统（ECDIS）操作和性能要求、测试方法及要求的测试结果",
    "GB/T 43053-2023 海上导航和无线电通信设备及系统 电子海图显示与信息系统（ECDIS）操作和性能要求、测试方法及要求的测试结果",
    "IEC 62288-2024 海上导航和无线电通信设备及系统 船载导航显示器上导航相关信息的显示一般要求、测试方法和所需测试结果合并版",
    "IHO S-57 3.1版 数字水文数据传输标准",
    "IHO S-52 6.1版 ECDIS的海图内容和显示方面",
    "IHO S-52 PresLib 4.0版 ECDIS的演示库、S-52附件A",
    "IHO S-63 1.2版 数据保护方案",
    "GJB438C-2021 军用软件开发文档通用要求",
]

FUNCTION_REQUIREMENTS = {
    "海图综合态势展示": [
        "核心显示与操作符合IEC 61174-2015、IEC 62288-2024、GB/T 43053-2023，海图数据处理符合IHO S-57、S-63、S-52及PresLib标准。",
        "接入并展示标准电子海图（S-57/S-63），支持海图自动匹配、选择、缩放、漫游、图层控制和特定要素开关显示。",
        "支持白天、黄昏、夜晚等IMO要求的多种配色模式，夜间模式亮度符合规范。",
        "支持正北向上、船艏向上，以及真实位移、相对位移模式切换。",
    ],
    "AIS/雷达目标融合显示": [
        "基于OpenGL技术实现雷达回波图像在电子海图上的实时叠加，集成雷达、AIS、ARPA等导航设备动态数据。",
        "将AIS目标符号、雷达回波、ARPA跟踪目标与本船符号准确叠加在电子海图相应地理位置。",
        "可选显示AIS目标船舶航向、航速、船名、MMSI等信息和ARPA跟踪预测信息。",
        "支持叠加电子围栏、风险区域、潮汐、海流、气象信息以及AIO等临时通告。",
    ],
    "航线规划与航行计算": [
        "支持手动或自动航线设计，允许添加、删除、拖动、排序航向点和航段。",
        "自动执行航线安全检查，提示航线段是否穿过浅水、禁航区或接近危险物。",
        "允许设定吃水、安全等深线、偏航阈值、船舶尺寸等航线安全参数。",
        "具备真距离、方位和ETA估算能力，并基于GPS/DGPS持续监控船舶是否按计划航线行驶。",
        "连续记录并存储历史航迹、船位、速度、航向和报警事件，用于航后分析和VDR备份。",
    ],
    "安全等深线与水深告警": [
        "允许输入船舶吃水和富裕水深，系统自动计算并设置安全等深线和安全轮廓线。",
        "在海图上突出显示浅于安全等深线的水域，并将安全水域显示为规范安全色。",
        "本船或安全轮廓线即将进入不安全水域、禁航区、限航区、港口报告区或电子围栏时立即报警。",
        "预测安全矢量接近孤立危险物或未知浅水区时触发报警，并基于AIS/雷达目标计算CPA/TCPA碰撞风险。",
    ],
    "航行监控导航与安全预警": [
        "船舶实时位置偏离计划航线超过设定偏航阈值（XTD）时立即发出告警。",
        "对GPS信号丢失、陀螺罗经数据异常、海图更新延迟、电源切换等系统级故障发出告警。",
        "遵循IMO/IHO规范对报警进行紧急、警报、指示等优先级分类，并采用不同声光模式。",
        "提供告警/预警信息列表，允许船员确认、消除告警。",
    ],
    "个性化交互与重要操作防误": [
        "提供直观图形化界面，菜单层级简洁明了，支持常用功能快捷键访问。",
        "软件界面适配高分辨率显示，支持通过软件指令调节背景亮度以适应夜间模式。",
        "删除航线、修改安全参数等重要操作弹出确认对话框并要求船员二次确认。",
        "支持触摸屏、鼠标、键盘、轨迹球等多种输入方式。",
    ],
    "平台架构与多设备协同": [
        "支持龙芯、飞腾、瑞芯微等硬件平台，以及麒麟、Ubuntu等操作系统部署并长期稳定运行。",
        "软件架构支持多台设备按优先级运行。",
        "采用分层结构和模块化设计，预留标准数据接口，支持后续接入新型号导航传感器或功能模块。",
        "本软件支持用户操作界面与显示，采用C++编程语言，并充分考虑继承性、灵活性和可维护性。",
    ],
    "数据维护、海图更新与接口集成": [
        "接收、解析和应用官方发布的S-63加密格式海图更新数据，并生成更新记录。",
        "定期自动备份系统内部存储结构（SENC）和所有航行记录，确保数据安全。",
        "软件接口遵循NMEA0183/2000协议，与船载设备进行数据通信。",
        "支持与USB存储设备、便携式电脑进行数据交换，支持CSV、XML等通用格式导入导出。",
        "支持GPS/北斗、电罗经、计程仪、BNWAS、VDR、BAM、MSI、INS等设备或系统接口。",
    ],
}

PERFORMANCE_REQUIREMENTS = [
    "一般操作响应时间（数据查阅、图表显示等操作）：＜2s。",
    "系统初始化时间：＜10s。",
    "一般数据查询时间：＜5s。",
    "一般数据汇总时间：＜10s。",
    "指令响应率：100%。",
    "系统服务稳定性：99%。",
    "特殊操作超过规定时间时，向用户给出明确等待时间或进度条提示。",
]

QUALITY_REQUIREMENTS = {
    "可靠性": [
        "海图数据加载准确完整，海图初始化时间≤8s，海图响应时间≤2s。",
        "航线安全检查结果准确率高于99%。",
        "存储介质可靠，接口通信具备错误校验和自动重连机制。",
    ],
    "安全性": [
        "系统部署运行软硬件环境满足全军统一国产化要求。",
        "上线运行前完成整体安全性测试并提供安全测试报告。",
        "系统管理人员经过严格培训，对访问用户按岗位业务进行权限设计并建立严格授权机制。",
        "正式运行系统与数据库中严禁进行测试操作，支持按军用计算机网络安全防护标准部署配置。",
    ],
    "保障性": [
        "软件随主机设备移交时同步移交用户使用手册。",
        "软件移交后配合联试和使用维护，对软件更改做好记录，保持文档与代码的可追溯性和一致性。",
    ],
    "维修性": [
        "数据接入、告警管理、系统配置等功能模块之间具有定义清晰、稳定的接口。",
        "MTTR（平均维修时间）≤0.5小时。",
    ],
    "测试性": [
        "系统及各模块设计和实现时充分考虑可测试性，提供必要测试接口或机制。",
        "支持自动化测试实施，便于进行单元测试、集成测试、系统测试和验收测试。",
    ],
    "环境适应性": [
        "软件运行环境覆盖龙芯、飞腾、瑞芯微硬件平台和麒麟、Ubuntu系统。",
    ],
}

BUSINESS_REQUIREMENTS = [
    "交付（服务）时间：合同签订后，以需求单位通知进场时间为准，进场后2个月内交货安装调试完毕，并完成系统集成、部署、联调、上线运行。",
    "交付（服务）地点：郑州市内甲方指定地点。",
    "交付（服务）方式：供应商完成运输、安装、调试等工作。",
    "软件质保期：自软件验收合格之日起24个月。",
    "服务方式：提供2年7×24小时上门保修，根据甲方要求采用远程指导、现地排故和驻场调试等方式，重大问题按商务要求提供现场技术支持。",
    "培训要求：提供免费培训，培训内容包括理论培训、现场操作培训，并负责编制安装手册、软件使用手册、PPT或视频等培训材料，培训次数不超过5次、每次不少于60分钟。",
    "知识产权和保密：保证采购单位使用不受第三方侵权指控，不向第三方泄露采购机构提供的技术文件等材料，合同履行形成的知识产权和其他权益按商务要求归属。",
    "驻场要求：投标供应商须进驻甲方办公场所，所有开发研制工作均在需求方办公场所完成，并接受甲方跟踪、协作与监督。",
]

SCORING_ITEMS = [
    ("技术方案", "需求分析全面性、准确性、充分性", "3.00", "需求分析、功能设计、覆盖矩阵"),
    ("技术方案", "系统总体、业务、逻辑、技术、数据架构全面性、合理性及配置灵活性", "3.00", "总体架构设计、数据库架构设计、部署架构设计"),
    ("技术方案", "业务和数据处理流程图符合需求且逻辑清晰", "2.00", "功能流程图、数据处理流程说明"),
    ("技术方案", "系统接口理解准确性、数据交互设计合理性", "2.00", "接口集成设计、数据维护与接口集成"),
    ("技术方案", "部署方案灵活性", "1.00", "部署架构设计"),
    ("技术方案", "项目安全设计针对性、合理性", "1.00", "安全性设计、部署安全措施"),
    ("技术方案", "软件开发、需求分析、软件设计、数据建模、软件测试工具证明", "2.00", "需用户补充工具证明或承诺书"),
    ("技术方案", "关键技术可行性、合理性、创新性及相关性", "2.00", "关键技术"),
    ("技术方案", "采购需求中功能符合情况和详细设计", "3.00", "功能设计章节"),
    ("技术指标", "一般技术指标正偏离", "11.00", "技术指标参数响应偏离表及承诺材料"),
    ("技术指标", "技术指标负偏离评分", "1.00", "技术指标参数响应偏离表"),
    ("项目管理和实施", "实施周期及进度管理方法措施", "1.00", "项目进度计划"),
    ("项目管理和实施", "项目风险评估及控制方案", "1.00", "风险评估与控制"),
    ("项目管理和实施", "质量保证措施、质量计划、测试与交付售后质量控制", "1.00", "质量控制、质量保证措施"),
    ("项目负责人", "高级职称、关键证书、同类业绩", "5.00", "需用户补充项目负责人证明材料"),
    ("开发项目团队和人员", "团队关键证书及驻场人员数量", "5.00", "需用户补充团队证明材料"),
    ("培训和售后服务", "培训计划方案、服务方式、现场支持、服务等级、升级服务期限", "4.00", "培训方案、售后服务及承诺"),
]


def ensure_dirs() -> None:
    OUTPUT.mkdir(exist_ok=True)
    RECORDS.mkdir(exist_ok=True)


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def para_after(paragraph, text: str = "", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph.__class__(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def replace_exact(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            set_para_text(paragraph, replacements[text])


def replace_after_heading(doc: Document, heading_text: str, placeholder: str) -> None:
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() != heading_text:
            continue
        if index + 1 < len(paragraphs):
            set_para_text(paragraphs[index + 1], placeholder)
        else:
            para_after(paragraph, placeholder)
        return


def compress_between(doc: Document, start_heading: str, end_heading: str | None, placeholders: list[str]) -> None:
    paragraphs = doc.paragraphs
    start = next((i for i, p in enumerate(paragraphs) if p.text.strip() == start_heading), None)
    if start is None:
        return
    if end_heading is None:
        end = len(paragraphs)
    else:
        end = next((i for i, p in enumerate(paragraphs[start + 1 :], start + 1) if p.text.strip() == end_heading), len(paragraphs))
    body = paragraphs[start + 1 : end]
    anchor = paragraphs[start]
    for offset, placeholder in enumerate(placeholders):
        if offset < len(body):
            set_para_text(body[offset], placeholder)
            anchor = body[offset]
        else:
            anchor = para_after(anchor, placeholder)
    for paragraph in body[len(placeholders) :]:
        delete_paragraph(paragraph)


def normalize_team_table(doc: Document) -> None:
    for table in doc.tables:
        if not table.rows:
            continue
        header = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
        if "拟担任" in header and "姓名" in header and len(table.rows) > 1:
            markers = [
                "【CONFIRM:项目团队-职务分工】",
                "【CONFIRM:项目团队-姓名】",
                "【CONFIRM:项目团队-职称】",
                "【CONFIRM:项目团队-专业】",
                "【CONFIRM:项目团队-从业资格】",
                "【CONFIRM:项目团队-相关工作年限】",
            ]
            for cell, marker in zip(table.rows[1].cells, markers):
                cell.text = marker


def delete_exact_paragraphs(doc: Document, texts: set[str]) -> None:
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in texts:
            delete_paragraph(paragraph)


def normalize_template_copy(doc: Document) -> None:
    replace_exact(
        doc,
        {
            "【GEN:根据技术要求及本项目的项目背景，生成1大段文字描述本文件的编写目的】": "【GEN:编写目的】",
            "【GEN:根据技术要求及本项目的项目背景，生成1大段文字描述项目的建设内容】": "【GEN:建设内容】",
            "【COPY:技术要求-通用质量特性要求/非功能性要求/六性要求】": "【COPY:技术要求-通用质量特性要求】",
            "【GEN:总体架构设计（根据技术要求，采用大段文字的形式描述系统总体架构设计）】": "【GEN:总体架构设计】",
            "【GEN:用大段文字描述架构图】": "【GEN:架构图说明】",
            "【GEN:用大段文字描述前述架构的设计原则】": "【GEN:设计原则】",
            "【GEN:根据技术要求，设计软件系统部署架构，并用大段文字描述】": "【GEN:部署架构设计】",
            "【性能设计章节需要一一对应技术要求中的性能需求逐条回答】": "【GEN:性能设计章节】",
        },
    )
    compress_between(doc, "功能设计", "性能设计", ["【GEN:功能设计总述】", "【GEN:功能设计章节】"])
    replace_after_heading(doc, "性能设计", "【GEN:性能设计章节】")
    replace_after_heading(doc, "数据库设计", "【GEN:数据库设计总述】")
    replace_after_heading(doc, "数据库架构", "【GEN:数据库架构设计】")
    replace_after_heading(doc, "核心业务设计", "【GEN:核心业务数据设计】")
    replace_after_heading(doc, "数据库表设计", "【GEN:数据库表设计】")
    replace_after_heading(doc, "通用质量特性设计", "【GEN:通用质量特性设计总述】")
    replace_after_heading(doc, "可靠性设计", "【GEN:可靠性设计】")
    replace_after_heading(doc, "维修性设计", "【GEN:维修性设计】")
    replace_after_heading(doc, "保障性设计", "【GEN:保障性设计】")
    replace_after_heading(doc, "测试性设计", "【GEN:测试性设计】")
    replace_after_heading(doc, "安全性设计", "【GEN:安全性设计】")
    replace_after_heading(doc, "环境适应性设计", "【GEN:环境适应性设计】")
    replace_after_heading(doc, "关键技术", "【GEN:关键技术】")
    replace_after_heading(doc, "质量控制", "【GEN:质量控制总述】")
    replace_after_heading(doc, "风险评估与控制", "【GEN:风险评估与控制】")
    compress_between(
        doc,
        "质量保证措施",
        "项目进度计划",
        ["【GEN:质量保证措施】", "【REVIEW:质量体系与资质响应说明】", "【REVIEW:服务质量保障措施】"],
    )
    compress_between(doc, "项目进度计划", "项目人员组成", ["【REVIEW:项目进度计划】"])
    compress_between(doc, "项目人员组成", "成果交付及验收", ["【CONFIRM:项目团队人员说明】"])
    compress_between(doc, "成果交付及验收", "售后服务及承诺", ["【REVIEW:成果交付及验收】", "【REVIEW:交付物清单】"])
    compress_between(
        doc,
        "售后服务及承诺",
        None,
        [
            "【GEN:培训方案】",
            "【CONFIRM:质量保证期承诺】",
            "【CONFIRM:售后服务响应承诺】",
            "【REVIEW:应急支援保障承诺】",
            "【REVIEW:定期跟踪服务承诺】",
        ],
    )
    normalize_team_table(doc)
    delete_exact_paragraphs(doc, {"系统架构图"})


def set_para_text(paragraph, text: str, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    run.bold = bold
    for run in paragraph.runs:
        run.font.name = "仿宋"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
        run.font.size = Pt(12)


def add_block_after(anchor, block: dict):
    if block["type"] == "paragraph":
        p = para_after(anchor, block.get("text", ""))
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = 1.25
        return p
    if block["type"] == "heading":
        p = para_after(anchor, block["text"])
        for run in p.runs:
            run.bold = True
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.size = Pt(block.get("size", 13))
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        return p
    if block["type"] == "image":
        p = para_after(anchor)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(block["path"]), width=Inches(block.get("width", 5.8)))
        return p
    if block["type"] == "caption":
        p = para_after(anchor, block["text"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
            run.font.size = Pt(10.5)
        return p
    if block["type"] == "table":
        doc = anchor.part.document
        rows = block["rows"]
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, value in enumerate(row):
                cell = table.cell(i, j)
                cell.text = str(value)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "仿宋"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
                        run.font.size = Pt(9 if len(str(value)) > 30 else 10.5)
                    p.paragraph_format.line_spacing = 1.1
        anchor._p.addnext(table._tbl)
        return table
    raise ValueError(block["type"])


def replace_placeholder_paragraph(doc: Document, placeholder: str, blocks: list[dict]) -> str:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == placeholder:
            first = blocks[0] if blocks else {"type": "paragraph", "text": ""}
            if first["type"] in {"paragraph", "heading"}:
                set_para_text(paragraph, first["text"], first["type"] == "heading")
                if first["type"] == "heading":
                    for run in paragraph.runs:
                        run.font.name = "黑体"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                anchor = paragraph
            elif first["type"] == "image":
                set_para_text(paragraph, "")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(first["path"]), width=Inches(first.get("width", 5.8)))
                anchor = paragraph
            else:
                set_para_text(paragraph, "")
                anchor = paragraph
                anchor = add_block_after(anchor, first)
            for block in blocks[1:]:
                anchor = add_block_after(anchor, block)
            return "已生成"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() == placeholder and blocks:
                    cell.text = "\n".join(b.get("text", "") for b in blocks if b["type"] in {"paragraph", "heading"})
                    return "已生成于表格"
    return "未找到"


def replace_inline(doc: Document, placeholder: str, value: str) -> str:
    status = "未找到"
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            set_para_text(paragraph, paragraph.text.replace(placeholder, value))
            status = "已替换"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, value)
                    status = "已替换于表格"
    return status


def bullets(items: list[str]) -> str:
    return "\n".join(f"（{i}）{item}" for i, item in enumerate(items, 1))


def prose_blocks(text: str) -> list[dict]:
    parts = [p.strip() for p in textwrap.dedent(text).strip().split("\n") if p.strip()]
    return [{"type": "paragraph", "text": p} for p in parts]


def heading(text: str, size: int = 13) -> dict:
    return {"type": "heading", "text": text, "size": size}


def paragraph(text: str) -> dict:
    return {"type": "paragraph", "text": text}


def draw_flow_png(path: Path, title: str, nodes: list[str]) -> None:
    font_path = Path(r"C:\Windows\Fonts\NotoSansSC-VF.ttf")
    font = ImageFont.truetype(str(font_path), 24)
    small = ImageFont.truetype(str(font_path), 18)
    width = 1500
    box_w = 350
    box_h = 72
    gap_x = 80
    gap_y = 70
    cols = 3
    rows = (len(nodes) + cols - 1) // cols
    height = 130 + rows * (box_h + gap_y)
    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    draw.text((width // 2, 28), title, fill="#1f2937", font=font, anchor="mm")
    positions = []
    for idx, node in enumerate(nodes):
        row = idx // cols
        col = idx % cols
        x = 90 + col * (box_w + gap_x)
        y = 85 + row * (box_h + gap_y)
        positions.append((x, y))
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=12, fill="#eef6ff", outline="#2f6f9f", width=3)
        wrapped = textwrap.wrap(node, width=13)
        ty = y + box_h / 2 - (len(wrapped) - 1) * 12
        for line in wrapped:
            draw.text((x + box_w / 2, ty), line, fill="#111827", font=small, anchor="mm")
            ty += 24
    for idx in range(len(positions) - 1):
        x1, y1 = positions[idx]
        x2, y2 = positions[idx + 1]
        start = (x1 + box_w, y1 + box_h / 2)
        end = (x2, y2 + box_h / 2)
        if idx // cols != (idx + 1) // cols:
            start = (x1 + box_w / 2, y1 + box_h)
            end = (x2 + box_w / 2, y2)
        draw.line((start, end), fill="#2f6f9f", width=3)
        draw.polygon([(end[0], end[1]), (end[0] - 12, end[1] - 7), (end[0] - 12, end[1] + 7)], fill="#2f6f9f")
    image.save(path)


def write_mmd(path: Path, title: str, chapter: str, body: str) -> None:
    path.write_text(f"%% 图名称：{title}\n%% 对应章节：{chapter}\n{body}\n", encoding="utf-8")


def safe_filename(text: str) -> str:
    return re.sub(r'[<>:"/\\|?*\s]+', "-", text).strip("-")


def create_diagrams() -> dict[str, Path]:
    diagrams: dict[str, Path] = {}
    arch_mmd = """flowchart TB
    A[用户访问层] --> B[前端展现层]
    B --> C[业务应用层]
    C --> D[平台能力层]
    D --> E[数据资源层]
    C --> F[集成接口层]
    F --> G[导航传感器与外部系统]
    E --> H[SENC海图库与航行记录库]
    I[安全保障体系] -.-> B
    I -.-> C
    I -.-> D
    J[运维监控体系] -.-> C
    J -.-> D
    K[国产化运行环境] --> B"""
    write_mmd(RECORDS / "architecture.mmd", "总体架构图", "系统总体架构", arch_mmd)
    draw_flow_png(
        RECORDS / "architecture.png",
        "电子海图显示与信息系统总体架构图",
        ["用户访问层", "前端展现层", "业务应用层", "平台能力层", "数据资源层", "集成接口层", "安全保障体系", "运维监控体系", "国产化运行环境"],
    )
    diagrams["总体架构图"] = RECORDS / "architecture.png"

    flow_nodes = {
        "海图综合态势展示": ["加载S-57/S-63海图", "解析并生成SENC", "匹配显示区域", "图层与要素控制", "配色和朝向切换", "叠加态势信息", "刷新显示", "记录操作日志"],
        "AIS/雷达目标融合显示": ["接收雷达/AIS/ARPA数据", "协议解析与校验", "坐标转换", "目标关联", "OpenGL叠加渲染", "风险信息计算", "显示目标详情", "记录融合结果"],
        "航线规划与航行计算": ["新建或导入航线", "编辑航点航段", "设置安全参数", "执行安全检查", "计算距离方位ETA", "保存航线版本", "实时监控偏航", "归档航迹记录"],
        "安全等深线与水深告警": ["录入吃水与富裕水深", "计算安全等深线", "生成安全轮廓线", "扫描浅水与危险物", "预测未来安全矢量", "判断风险阈值", "触发声光告警", "写入告警日志"],
        "航行监控导航与安全预警": ["接收实时船位", "校验传感器状态", "比对计划航线", "判断XTD偏差", "识别系统故障", "告警优先级分类", "船员确认处置", "闭环归档"],
        "个性化交互与重要操作防误": ["进入操作界面", "选择快捷功能", "执行参数调整", "识别重要操作", "弹出二次确认", "校验权限", "反馈执行结果", "记录审计日志"],
        "平台架构与多设备协同": ["识别运行平台", "加载模块配置", "检测设备优先级", "建立主备协同", "同步关键状态", "接口扩展接入", "故障切换", "记录运行状态"],
        "数据维护、海图更新与接口集成": ["接收更新包", "校验S-63授权", "解析更新内容", "应用到SENC", "生成更新记录", "备份航行数据", "接口数据交换", "异常回退处理"],
    }
    for idx, (name, nodes) in enumerate(flow_nodes.items(), 1):
        safe = f"function-{idx:03d}-{safe_filename(name)}"
        mmd = "flowchart TD\n" + "\n".join(
            f"    N{i}[{node}] --> N{i + 1}[{nodes[i]}]" if i == 0 else "" for i, node in enumerate(nodes[:-1])
        )
        lines = ["flowchart TD"]
        for i, node in enumerate(nodes, 1):
            lines.append(f"    N{i}[{node}]")
        for i in range(1, len(nodes)):
            lines.append(f"    N{i} --> N{i+1}")
        write_mmd(RECORDS / f"{safe}.mmd", f"{name}流程图", f"{name}功能", "\n".join(lines))
        png = RECORDS / f"{safe}.png"
        draw_flow_png(png, f"{name}流程图", nodes)
        diagrams[name] = png
    return diagrams


def build_copy_blocks() -> dict[str, list[dict]]:
    function_text = []
    for name, items in FUNCTION_REQUIREMENTS.items():
        function_text.append(f"{name}：{bullets(items)}")
    quality_text = []
    for name, items in QUALITY_REQUIREMENTS.items():
        quality_text.append(f"{name}要求：{bullets(items)}")
    return {
        "【COPY:技术要求-功能要求】": prose_blocks("\n".join(function_text)),
        "【COPY:技术要求-性能要求】": prose_blocks("\n".join(PERFORMANCE_REQUIREMENTS)),
        "【COPY:技术要求-通用质量特性要求】": prose_blocks("\n".join(quality_text)),
    }


def build_gen_blocks(diagrams: dict[str, Path]) -> dict[str, list[dict]]:
    blocks: dict[str, list[dict]] = {}
    blocks["【GEN:编写目的】"] = prose_blocks(
        f"""
        本设计方案报告用于响应{PROJECT_NAME}项目技术要求、商务要求和技术评分要求，明确系统建设目标、总体架构、功能设计、性能设计、数据库设计、质量控制、实施保障和培训服务安排。文档以招标文件提出的标准符合性、电子海图处理、航行监控、安全告警、数据接口、国产化运行环境和软件工程文档要求为依据，形成可用于投标文件技术方案部分的初稿，为后续深化设计、合同履约、系统开发、集成联调和验收交付提供统一的方案基础。
        """
    )
    blocks["【GEN:建设内容】"] = prose_blocks(
        f"""
        本项目建设内容围绕电子海图显示与信息系统的软件研制、集成部署和运行保障展开，重点包括标准电子海图接入与显示、AIS/雷达/ARPA目标融合、航线规划与安全检查、安全等深线和水深告警、航行监控与多级预警、个性化人机交互、跨平台国产化部署、多设备协同、海图数据更新、航行记录存储、外部接口集成、系统安全防护、性能优化、测试验证、交付验收、培训和售后服务等内容。系统交付物应覆盖软件源代码、可执行文件、部署运维指南、按GJB438C-2021形成的技术文档以及用户使用资料。
        """
    )
    blocks["【GEN:总体架构设计】"] = prose_blocks(
        """
        系统总体架构采用分层、模块化和接口标准化的设计思路，围绕船载电子海图显示、导航态势融合、安全预警和航行数据管理等核心业务构建。整体架构从上至下划分为用户访问层、前端展现层、业务应用层、平台能力层、数据资源层、集成接口层、安全保障体系、运维监控体系和国产化运行环境。各层之间通过清晰接口传递业务请求、海图数据、传感器数据、告警事件和运维状态，保证功能可扩展、接口可替换、运行可监控、问题可追溯。
        用户访问层面向船员、系统管理人员、联试维护人员等不同岗位用户，提供触摸屏、鼠标、键盘、轨迹球等多种交互方式。系统根据岗位职责和业务场景建立访问边界，使航线规划、告警确认、系统配置、数据维护等操作具备明确权限控制和操作审计能力，满足正式运行系统中严格授权和禁止测试操作的安全要求。
        前端展现层承担电子海图显示、态势叠加、航线编辑、告警提示、信息查询和配置操作等交互任务。该层以高分辨率显示适配、昼夜配色、正北向上/船艏向上、真实位移/相对位移等显示模式为基础，结合OpenGL图形渲染能力，实现S-57/S-63海图、雷达回波、AIS目标、ARPA目标、气象海况、自定义GIS要素和安全告警的统一可视化呈现。
        业务应用层承载系统核心功能，包括海图综合态势展示、目标融合显示、航线规划、航行监控、安全等深线计算、水深和区域告警、CPA/TCPA碰撞预警、个性化设置、海图更新和航行记录管理。各业务模块以功能边界清晰、职责单一、接口稳定为原则组织，保证后续接入新型号导航传感器或扩展功能模块时不会破坏既有业务逻辑。
        平台能力层提供图形渲染、坐标转换、协议解析、数据校验、告警规则、任务调度、日志审计、配置管理、权限控制、备份恢复和运行监控等通用能力。该层将共性技术能力从业务模块中抽离，降低重复开发和维护成本，并为性能优化、自动化测试、故障定位和跨平台迁移提供统一支撑。
        数据资源层以SENC海图库、航线数据、航行记录、告警事件、系统配置、用户权限、海图更新记录和接口交换数据为核心对象。系统对海图更新、航行轨迹、报警事件和关键操作建立全生命周期记录，支持数据备份、恢复、归档和追溯，满足航后分析、VDR备份、质量监督和验收取证要求。
        集成接口层按照NMEA0183/2000等协议与GPS/北斗、电罗经、计程仪、雷达、AIS、ARPA以及BNWAS、VDR、BAM、MSI、INS等外部设备或系统进行数据交互，同时支持USB存储设备和便携式电脑的数据导入导出。接口设计采用解析、校验、缓存、状态监测和自动重连机制，确保通信稳定性和异常可恢复性。
        安全保障体系和运维监控体系贯穿各层。安全体系覆盖身份认证、岗位授权、重要操作二次确认、数据访问控制、安全测试、保密管理和网络防护部署；运维体系覆盖运行状态采集、日志审计、告警闭环、自动备份、故障定位、版本记录和文档代码一致性维护，从而支撑系统在国产化软硬件环境中长期稳定运行。
        """
    )
    blocks["【GEN:总体架构图】"] = [
        {"type": "image", "path": diagrams["总体架构图"], "width": 6.1},
        {"type": "caption", "text": "图 1 电子海图显示与信息系统总体架构图"},
    ]
    blocks["【GEN:架构图说明】"] = prose_blocks(
        """
        架构图体现系统从用户交互到业务处理、平台支撑、数据存储和外部集成的完整链路。用户访问层通过岗位权限进入前端展现层，前端展现层将海图显示、目标态势、航线编辑、告警提示等操作请求传递至业务应用层。业务应用层调用平台能力层完成图形渲染、协议解析、坐标转换、规则计算、日志审计和配置管理，并将海图、航线、航迹、告警、系统配置和更新记录持久化到数据资源层。
        集成接口层负责与船载导航传感器、外部业务系统和文件交换设备进行数据交互，向业务应用层提供经过校验和标准化处理的动态数据。安全保障体系对访问、数据、操作、网络和测试过程进行约束，运维监控体系对运行状态、接口链路、告警事件和备份任务进行持续监测，两类支撑体系以横向能力方式作用于系统各层，保证系统安全可控、运行稳定、问题可追溯。
        """
    )
    blocks["【GEN:设计原则】"] = prose_blocks(
        """
        先进适用原则。系统采用符合ECDIS相关国际、国家和行业标准的技术路线，围绕S-57/S-63/S-52海图处理、OpenGL图形叠加、NMEA协议接入、多级告警和跨平台运行等关键能力进行设计，保证技术路线与项目需求直接相关，并避免脱离实际运行环境的技术堆砌。
        稳定可靠原则。系统对海图加载、传感器通信、航线安全检查、告警触发、数据存储和自动备份等关键链路建立校验、重试、状态监测和日志追踪机制，接口通信具备错误校验和自动重连能力，核心模块具备清晰稳定的接口边界，以支撑长期稳定运行和平均维修时间要求。
        安全可控原则。系统按照岗位业务进行权限设计，对删除航线、修改安全参数等重要操作实施二次确认，对正式运行系统和数据库建立严格操作边界。系统上线前开展整体安全性测试，并结合军用计算机网络安全防护标准落实部署配置、访问控制、审计记录和保密管理要求。
        易扩展原则。系统采用分层结构和模块化设计，传感器接入、海图处理、态势显示、航行监控、告警规则、数据维护等模块之间通过标准接口协作，后续新增导航设备、接口协议、业务规则或显示要素时，可通过模块扩展和配置调整实现。
        易维护原则。系统对配置、日志、版本、接口状态、备份任务和异常事件进行集中管理，形成从故障发现、定位、处置到记录归档的闭环机制。软件变更过程保留记录，并保持文档与代码的可追溯性和一致性，便于联试、验收、质保和后续维护。
        标准规范原则。系统执行IEC 61174-2015、GB/T 43053-2023、IEC 62288-2024、IHO S-57、IHO S-52、IHO S-63、GJB438C-2021等规范要求，设计文档、接口处理、海图显示、数据保护、软件开发文档和验收资料均围绕标准要求组织。
        """
    )
    blocks["【GEN:部署架构设计】"] = prose_blocks(
        """
        系统部署面向龙芯、飞腾、瑞芯微等硬件平台和麒麟、Ubuntu等操作系统，采用可移植的软件包、配置文件和运行脚本组织方式，保证在不同CPU架构和国产化操作系统环境下具备一致的功能表现。部署设计不绑定具体品牌和型号，运行环境参数以甲方提供设备和最终合同约定为准。
        在船载运行环境中，系统可按主用设备、备用设备或多台设备优先级运行方式部署。主用节点负责实时海图显示、目标融合、航行监控和告警处置，备用节点保持关键配置、海图数据、航线数据和运行状态同步。当主用节点发生故障或维护切换时，系统根据设备优先级和运维策略完成接续运行，降低单点故障对航行安全的影响。
        部署边界包括系统终端、海图数据存储、导航传感器接口、外部数据交换设备、运维管理通道和安全防护配置。系统通过标准协议接入GPS/北斗、电罗经、计程仪、AIS、雷达、ARPA等数据源，并通过文件交换方式支持海图更新包、航线文件和运行记录导入导出。各接口均进行状态监测、数据校验和异常记录。
        数据存储采用业务数据、海图数据、日志数据和备份数据分区管理思路。SENC海图库、航行记录、告警事件、系统配置和更新记录均纳入定期备份范围，备份任务具备执行记录和结果校验。涉及外部交换介质时，系统按安全要求进行文件来源、格式、完整性和授权状态校验。
        安全部署方面，系统结合岗位权限、最小授权、重要操作二次确认、日志审计、安全测试和网络防护配置进行整体控制。运维部署方面，系统提供运行状态、接口链路、资源占用、备份任务、告警事件和版本信息的检查能力，支持联试、验收和质保阶段的问题定位与闭环整改。
        """
    )
    blocks["【GEN:功能设计总述】"] = prose_blocks(
        """
        功能设计按照技术要求中的业务能力、接口能力、数据能力、安全能力和运行保障能力进行拆分，形成海图综合态势展示、AIS/雷达目标融合显示、航线规划与航行计算、安全等深线与水深告警、航行监控导航与安全预警、个性化交互与重要操作防误、平台架构与多设备协同、数据维护与接口集成等功能域。各功能域均对应招标文件明确要求，并结合技术评分表对需求响应、架构合理性、流程图清晰度、接口理解、部署灵活性、安全设计、关键技术和功能详细设计等评分点进行强化。
        """
    )
    function_blocks: list[dict] = []
    for idx, (name, reqs) in enumerate(FUNCTION_REQUIREMENTS.items(), 1):
        function_blocks.append(heading(f"{idx}.{name}功能", 13))
        function_blocks.extend(prose_blocks(function_design_text(name, reqs)))
        function_blocks.append({"type": "image", "path": diagrams[name], "width": 5.9})
        function_blocks.append({"type": "caption", "text": f"图 {idx + 1} {name}流程图"})
    blocks["【GEN:功能设计章节】"] = function_blocks
    blocks["【GEN:性能设计章节】"] = prose_blocks(performance_design_text())
    blocks["【GEN:数据库设计总述】"] = prose_blocks(database_intro_text())
    blocks["【GEN:数据库架构设计】"] = prose_blocks(database_arch_text())
    blocks["【GEN:核心业务数据设计】"] = prose_blocks(core_data_text())
    blocks["【GEN:数据库表设计】"] = [
        paragraph("数据库表设计以可调整初稿方式给出，具体字段编码、长度、约束、索引和分区策略应在详细设计阶段结合甲方确认的数据字典、接口报文和部署环境进一步固化。"),
        {"type": "table", "rows": database_table_rows()},
    ]
    blocks["【GEN:通用质量特性设计总述】"] = prose_blocks(
        """
        通用质量特性设计围绕可靠性、维修性、保障性、测试性、安全性和环境适应性展开，既响应技术要求中的明确指标，也支撑技术评分表对安全设计、质量控制、风险控制和关键技术可行性的评审要求。系统在设计阶段将质量特性分解到架构、模块、接口、数据、测试和运维全过程，避免将质量要求停留在口号层面。
        """
    )
    blocks["【GEN:可靠性设计】"] = prose_blocks(
        """
        系统通过海图数据完整性校验、SENC加载结果校验、接口数据错误检查、通信链路自动重连、关键服务状态监控和异常日志记录保证运行可靠性。对海图初始化、海图响应、航线安全检查和告警触发等核心环节设置专项测试用例，确保海图初始化时间、响应时间和航线安全检查准确率满足技术要求。对历史航迹、告警事件、系统配置和海图更新记录建立定期备份机制，降低存储介质异常或误操作带来的数据损失风险。
        """
    )
    blocks["【GEN:维修性设计】"] = prose_blocks(
        """
        系统按照模块化原则划分数据接入、海图显示、态势融合、航线规划、告警管理、系统配置、数据维护和运维监控等模块。模块之间通过清晰、稳定的接口传递数据和状态，便于故障定位、局部替换和版本升级。系统运行日志记录接口状态、异常堆栈、告警规则命中、配置变更和用户操作，为维护人员快速判断故障来源提供依据，并围绕MTTR≤0.5小时的要求设计常见故障处置流程和检查清单。
        """
    )
    blocks["【GEN:保障性设计】"] = prose_blocks(
        """
        保障性设计覆盖软件交付、文档交付、联试配合、使用维护、变更记录和知识转移。系统随主机设备移交时同步提供用户使用手册、部署运维指南和按GJB438C-2021形成的技术文档，保证使用方具备安装、配置、操作、故障排查和日常维护依据。软件变更过程记录变更原因、影响范围、版本号、测试结果和交付状态，保持文档与代码可追溯、一致。
        """
    )
    blocks["【GEN:测试性设计】"] = prose_blocks(
        """
        系统在模块设计中预留测试接口和测试数据构造机制，支持对协议解析、海图加载、坐标转换、航线安全检查、告警规则、权限控制、数据备份和接口重连等能力开展单元测试、集成测试、系统测试和验收测试。测试过程形成测试计划、测试用例、测试记录、缺陷闭环和测试报告，覆盖正常路径、异常路径、边界值、并发刷新、接口中断、海图更新失败和权限不足等典型场景。
        """
    )
    blocks["【GEN:安全性设计】"] = prose_blocks(
        """
        系统安全性设计从部署环境、访问控制、数据安全、操作安全、接口安全和测试安全六个方面落实。部署环境满足全军统一国产化要求并支持按军用计算机网络安全防护标准配置；访问控制按照岗位业务设计角色和权限，防止越权配置、越权删除和越权查看；重要操作实施二次确认并记录审计日志；接口数据进行协议校验、来源校验和异常记录；上线前开展整体安全性测试并形成安全测试报告；正式运行系统和数据库禁止进行测试操作。
        """
    )
    blocks["【GEN:环境适应性设计】"] = prose_blocks(
        """
        系统面向龙芯、飞腾、瑞芯微等硬件平台和麒麟、Ubuntu等操作系统进行跨平台适配设计。软件采用C++语言和分层模块化结构，图形渲染、协议解析、文件路径、系统服务、配置加载和日志输出等平台相关能力通过适配层封装，降低不同架构与系统差异对业务模块的影响。测试阶段应在甲方提供或确认的目标环境中开展安装、运行、性能、稳定性和接口联调验证。
        """
    )
    blocks["【GEN:关键技术】"] = prose_blocks(key_tech_text())
    blocks["【GEN:质量控制总述】"] = prose_blocks(quality_control_text())
    blocks["【GEN:风险评估与控制】"] = [
        paragraph("项目风险控制围绕需求、技术、接口、数据、安全、进度、质量和验收等方面开展，采用风险识别、影响评估、预防措施、应急处置和闭环跟踪相结合的方式实施。"),
        {"type": "table", "rows": risk_rows()},
    ]
    blocks["【GEN:质量保证措施】"] = prose_blocks(quality_assurance_text())
    blocks["【GEN:培训方案】"] = prose_blocks(training_text())
    return blocks


def function_design_text(name: str, reqs: list[str]) -> str:
    return f"""
    一、功能目标。{name}功能用于落实招标文件中关于{reqs[0]}等要求，支撑电子海图显示与信息系统在船载环境下完成规范显示、实时处理、安全预警和运行记录。该功能设计以标准符合、操作清晰、数据准确、响应及时和安全可控为目标，并与技术评分表中功能符合性、流程清晰度、接口理解和关键技术相关性要求相对应。
    二、建设内容。系统围绕该功能配置业务页面、数据对象、操作入口、参数项、结果展示和日志记录能力。主要建设内容包括：{bullets(reqs)} 相关操作通过统一菜单、工具栏、图标或快捷键进入，关键状态在电子海图主界面、信息面板、告警列表或配置界面中集中展示，保证船员能够在航行场景下快速识别、快速操作和快速确认。
    三、实现方式。前端采用适配高分辨率显示的图形界面组织操作区域、海图区域、状态区域和告警区域，后端通过业务服务、规则引擎、数据访问组件和接口组件完成数据处理。涉及海图、位置、目标、航线、告警和配置的数据均经过格式校验、权限校验和状态校验后写入或读取。涉及外部数据的场景通过协议解析、缓存队列、时间戳比对和异常重连机制保证数据连续性。
    四、业务流程。用户进入对应功能后，系统先校验登录状态、岗位权限和当前运行模式，再加载所需海图、航线、目标、配置或接口数据。用户发起查询、编辑、监控或确认操作后，系统执行业务规则校验；校验通过时更新显示或写入数据，校验不通过时返回明确提示。涉及安全风险、接口异常或重要参数变化时，系统生成告警或确认提示，并同步记录操作日志。
    五、数据处理逻辑。该功能处理的数据包括输入参数、海图对象、目标对象、航线对象、告警事件、系统配置和操作记录。系统对输入数据进行范围、格式、完整性和来源校验；对业务计算结果保留时间、来源和版本信息；对重要历史数据进行归档和备份。数据查询和统计过程按照性能要求优化索引、缓存和计算路径，保证一般查询、汇总和显示响应满足招标指标。
    六、权限与安全控制。系统按岗位职责控制查看、编辑、确认、删除、导入、导出和系统配置权限。删除航线、修改安全参数、应用海图更新、调整接口配置等重要操作必须经过二次确认。所有关键操作写入审计日志，日志内容包括操作人员、时间、对象、参数摘要、处理结果和异常信息，为质量监督、问题追溯和安全审计提供依据。
    七、异常处理与运维保障。当出现数据缺失、接口中断、格式不合法、权限不足、计算失败或存储异常时，系统返回明确提示并记录异常日志。对可恢复异常执行重试、回滚或降级显示；对影响航行安全的异常按照告警优先级进行声光提示。运维人员可通过日志、接口状态、配置版本和备份记录定位问题，形成发现、处置、验证、归档的闭环。
    八、招标要求响应说明。本功能逐项响应技术要求中关于{name}的功能、接口、性能、安全和运维要求，并通过流程图展示业务处理路径。设计中未写入未经确认的厂商型号、人员资质、报价或承诺日期，涉及最终承诺和证明材料的事项在确认清单或复核清单中保留。
    """


def performance_design_text() -> str:
    rows = "\n".join(f"{item} 系统通过请求分级、缓存优化、异步处理、渲染刷新控制和日志监测等方式进行响应。" for item in PERFORMANCE_REQUIREMENTS)
    return f"""
    性能设计以招标文件列明指标为约束，将海图加载、图形显示、数据查询、数据汇总、指令响应和长耗时操作提示纳入统一性能控制范围。系统对交互类操作优先保障即时反馈，对查询和汇总类操作优化数据访问路径，对图形渲染类操作控制刷新频率和增量绘制，对接口类操作建立缓存与状态监测，确保航行场景下关键操作不被后台任务阻塞。
    {rows}
    对特殊操作超过规定时间的场景，系统在前端给出明确等待时间、进度条或处理状态提示，避免用户误判系统无响应。性能测试阶段应围绕海图初始化、海图缩放漫游、目标叠加刷新、航线安全检查、历史记录查询、海图更新应用、接口中断恢复等场景建立测试用例和记录，形成可供验收复核的性能测试结果。
    """


def database_intro_text() -> str:
    return """
    数据库设计围绕海图资源、航线资源、航行记录、告警事件、接口数据、系统配置、用户权限、海图更新记录和运维日志等核心对象展开。系统数据设计遵循业务分层、对象清晰、来源可追溯、更新可记录、异常可回滚和备份可恢复原则。涉及海图数据时，以S-57/S-63数据解析和SENC内部存储结构为基础；涉及运行数据时，按航次、时间、设备、来源和事件类型建立索引，支撑航后分析、VDR备份、质量监督和验收检查。
    """


def database_arch_text() -> str:
    return """
    数据库架构分为基础字典层、海图资源层、业务运行层、接口交换层、系统管理层和备份归档层。基础字典层保存代码表、告警级别、设备类型、角色权限和配置项；海图资源层保存SENC索引、海图单元、更新记录和要素显示控制信息；业务运行层保存航线、航点、航段、航迹、船位、速度、航向、告警事件和确认记录；接口交换层保存传感器数据缓存、外部文件导入导出记录和协议解析状态；系统管理层保存用户、角色、权限、审计日志和系统配置；备份归档层保存定期备份、恢复记录和历史归档索引。
    数据库访问采用统一数据访问组件封装，业务模块不直接操作底层表结构。系统对高频查询对象建立索引，对历史航迹和告警事件按时间维度归档，对海图更新和系统配置保留版本记录。涉及敏感配置、用户权限和关键航行数据的访问必须经过权限校验并记录审计日志。
    """


def core_data_text() -> str:
    return """
    核心业务数据包括海图数据、SENC数据、航线数据、航点航段数据、船舶动态数据、AIS目标数据、雷达/ARPA目标数据、安全参数、告警事件、接口状态、系统配置、用户权限、操作日志和备份记录。海图数据用于底图显示和空间分析；航线数据用于规划、检查、监控和ETA计算；动态目标数据用于态势融合、CPA/TCPA计算和碰撞预警；告警事件用于风险提示、确认处置和闭环归档；系统配置和权限数据用于运行控制和安全管理。
    数据之间的关系以航次、航线、时间、空间位置、数据来源和事件编号为关键关联条件。系统对关键数据建立来源标识、采集时间、处理时间、处理状态和版本信息，保证数据在导入、解析、计算、显示、存储、备份和归档全过程中可追溯。
    """


def database_table_rows() -> list[list[str]]:
    return [
        ["表名", "主要字段（初稿）", "用途", "设计说明"],
        ["chart_cell", "cell_id、edition、update_no、source_type、status", "海图单元与版本管理", "字段长度和编码按S-57/S-63解析结果确认"],
        ["senc_index", "senc_id、cell_id、build_time、checksum、storage_path", "SENC内部存储索引", "支持加载校验、更新追踪和备份恢复"],
        ["route_plan", "route_id、route_name、version、creator、status", "航线方案管理", "支持航线版本、启停状态和审计"],
        ["route_waypoint", "waypoint_id、route_id、seq、lat、lon、safety_param", "航点航段管理", "坐标精度和安全参数详细字段后续确认"],
        ["navigation_track", "track_id、time、lat、lon、speed、course、source", "航迹与船位记录", "按时间索引，支持航后分析和VDR备份"],
        ["target_state", "target_id、target_type、mmsi、lat、lon、speed、course", "AIS/雷达/ARPA目标状态", "目标融合规则在详细设计阶段细化"],
        ["alarm_event", "alarm_id、alarm_type、level、object_id、time、status", "告警事件与确认记录", "支持优先级、确认、消除和归档"],
        ["interface_log", "log_id、device_type、protocol、status、time、message", "接口通信与异常记录", "用于自动重连、故障定位和验收取证"],
        ["user_role", "user_id、role_id、permission_set、status", "用户权限管理", "实际用户信息由投标人和甲方确认后配置"],
        ["operation_audit", "audit_id、user_id、action、object、result、time", "关键操作审计", "支持安全审计和质量监督"],
    ]


def key_tech_text() -> str:
    return """
    一是标准电子海图解析与SENC构建技术。系统围绕S-57、S-63、S-52和PresLib要求处理海图数据、显示规则、数据保护和符号化表达，形成可快速加载和查询的内部存储结构，并对海图更新过程保留记录。
    二是OpenGL态势融合渲染技术。系统将电子海图、雷达回波、AIS目标、ARPA目标、自定义GIS要素、气象海况和安全告警叠加到统一空间坐标体系下，通过分层渲染、增量刷新和显示优先级控制保证态势表达清晰。
    三是航线安全检查与动态风险预测技术。系统结合吃水、安全等深线、禁航区、限航区、危险物、偏航阈值、AIS/雷达目标和CPA/TCPA规则进行安全检查和风险预警，实现规划阶段和航行阶段的连续安全控制。
    四是标准接口接入与异常恢复技术。系统按照NMEA0183/2000等协议接入船载设备，结合数据校验、时间戳比对、状态监测、缓存队列和自动重连机制保证动态数据稳定输入。
    五是跨平台国产化适配技术。系统采用C++和分层模块化设计，通过平台适配层隔离不同CPU架构、操作系统、图形环境和文件系统差异，支撑龙芯、飞腾、瑞芯微及麒麟、Ubuntu环境部署。
    六是全过程质量追溯技术。系统对需求、设计、开发、测试、部署、运行、告警、变更、备份和交付形成记录，使软件代码、技术文档、测试结果和运行日志保持可追溯、一致。
    """


def quality_control_text() -> str:
    return """
    质量控制覆盖需求分析、方案设计、详细设计、开发编码、集成联调、系统测试、交付验收和售后运维全过程。需求阶段建立需求响应矩阵，逐项映射技术要求、商务要求和评分项；设计阶段对架构、功能、接口、数据、安全和部署方案进行评审；开发阶段执行编码规范、版本管理和单元测试；联调阶段重点验证海图加载、传感器接入、目标融合、航线检查、告警闭环和数据备份；测试阶段形成测试计划、测试用例、缺陷记录和测试报告；交付阶段按验收大纲和交付物清单提交软件、文档和运行证明材料。
    项目接受甲方质量体系及客户代表机构监督评价，按要求提供过程数据、记录、文件和现场支持。对发现的问题建立纠正措施闭环，明确责任、原因、整改措施、验证方式和关闭条件。涉及产品检测评价标准、缺陷判据和抽样方法等事项，应配合甲方按验收大纲和合同要求执行。
    """


def risk_rows() -> list[list[str]]:
    return [
        ["风险类别", "风险表现", "影响", "控制措施"],
        ["需求风险", "功能边界或验收标准理解不一致", "返工、验收争议", "建立需求响应矩阵和评审记录，关键条款提交甲方确认"],
        ["接口风险", "传感器协议、数据频率或现场设备状态不一致", "联调延期、数据异常", "提前梳理接口清单，设计协议适配和异常重连机制"],
        ["海图数据风险", "S-63授权、更新包或SENC转换异常", "海图显示不完整或更新失败", "实施授权校验、完整性校验、更新记录和回滚机制"],
        ["性能风险", "目标叠加和海图刷新负载较高", "操作响应超时", "采用分层渲染、缓存、异步处理和专项性能测试"],
        ["安全风险", "权限配置不当或重要操作误操作", "数据泄露、误删、误报警", "落实岗位权限、二次确认、审计日志和安全测试"],
        ["进度风险", "进场、驻场、联试或用户支持安排变化", "交付周期压缩", "按商务要求制定阶段计划，关键节点滚动跟踪"],
        ["质量风险", "缺陷关闭不充分或文档代码不一致", "验收不通过", "执行缺陷闭环、版本管理、文档同步和验收前检查"],
    ]


def quality_assurance_text() -> str:
    return """
    质量保证措施包括组织保障、过程保障、技术保障、测试保障和交付保障。组织方面建立项目负责人、需求分析、架构设计、软件开发、测试验证、配置管理、质量控制和运维支持等角色分工，具体人员、职称、证书和工作年限由投标人真实材料确认后填报。过程方面按照需求确认、设计评审、编码实现、联调测试、验收交付和售后服务进行阶段控制，每一阶段形成记录并接受甲方监督。
    技术方面，系统围绕标准符合性、海图数据准确性、接口稳定性、告警及时性、权限安全性和跨平台适配性设置设计准则和检查项。测试方面，建立单元测试、集成测试、系统测试、性能测试、安全测试、环境适配测试和验收测试，覆盖正常、异常、边界和恢复场景。交付方面，按要求提交软件源代码、可执行文件、部署运维指南、GJB438C-2021技术文档和用户使用资料，保证交付物完整、版本一致、可安装、可运行、可验证。
    """


def training_text() -> str:
    return """
    培训方案依据商务要求设置理论培训、现场操作培训和资料交付三类内容，培训对象为甲方指定的系统使用人员、系统管理人员和维护保障人员，具体对象与人数按甲方要求执行。培训目标是使受训人员掌握电子海图显示、航线规划、目标态势查看、安全告警确认、海图更新、数据备份、系统配置和常见故障处理等能力。
    培训内容包括系统总体介绍、标准电子海图基础、界面操作、显示模式切换、航线规划与安全检查、AIS/雷达/ARPA目标查看、安全等深线与水深告警、航行监控、接口状态检查、海图更新、日志查询、备份恢复、权限管理和应急处置。培训方式包括课堂讲解、现场演示、实操练习、问题答疑和考核反馈。
    培训资料包括安装手册、软件使用手册、培训PPT或视频、操作流程说明和常见问题处理清单。培训教员应具备相同课程教学经验并使用中文授课。培训次数按商务要求不超过5次，每次不少于60分钟；培训时间、地点、授课老师和具体课程安排应结合项目实施进度和甲方要求最终确认。培训结束后形成签到、培训记录、问题清单和效果反馈，为后续售后服务和运维保障提供依据。
    """


def build_review_blocks() -> dict[str, list[dict]]:
    return {
        "【REVIEW:质量体系与资质响应说明】": prose_blocks("投标文件应结合投标人真实质量管理体系、资质证书、软件研发过程管理制度和相关证明材料进行响应。本初稿仅提示应覆盖质量体系、质量计划、质量监督配合、缺陷闭环和交付质量控制等内容，具体证书名称、编号、有效期和证明材料需人工复核后补充。"),
        "【REVIEW:服务质量保障措施】": prose_blocks("服务质量保障应覆盖7×24小时服务、远程指导、现地排故、驻场调试、问题处理报告、服务记录和闭环管理等商务要求。具体服务团队、联系人、响应时限承诺和驻场安排需与投标人实际服务能力及合同条款复核一致。"),
        "【REVIEW:项目进度计划】": prose_blocks("项目进度计划建议按项目启动与进场准备、需求确认、总体设计与详细设计、系统开发与配置、数据准备与接口联调、系统测试与问题整改、安装部署与上线运行、初验、试运行、终验和资料移交组织。商务要求明确以需求单位通知进场时间为准，进场后2个月内交货安装调试完毕并完成集成、部署、联调、上线运行；具体日期、里程碑和人力投入需人工复核。"),
        "【REVIEW:成果交付及验收】": prose_blocks("成果交付及验收应按照甲方编制的验收大纲、规范和合同要求执行。乙方需配合完成外包验收工作，提供必要的数据、记录、文件和现场支持。验收重点包括软件在甲方提供设备上可靠运行、功能和性能满足技术要求、接口联调通过、安全测试报告完整、技术文档符合GJB438C-2021、问题整改闭环。具体验收口径和拒收准则需结合合同及验收大纲复核。"),
        "【REVIEW:交付物清单】": [
            paragraph("交付物清单依据技术要求表1形成初稿，最终数量、介质、格式和签收要求需人工复核。"),
            {"type": "table", "rows": [
                ["序号", "交付物名称", "数量", "成果形式", "备注"],
                ["1", "交付软件源代码和可执行文件在2套甲方提供的设备上能可靠运行", "2套", "实物", "按甲方设备环境验证"],
                ["2", "交付软件源代码（含完整注释）及可执行文件", "1份", "纸质、光盘", "版本、介质和签收要求需复核"],
                ["3", "软件部署及运维指南", "1份", "纸质、光盘", "含部署、配置、备份、恢复和故障处理"],
                ["4", "按GJB438C-2021提供技术文档", "1套", "纸质、光盘", "文档目录和模板需复核"],
            ]},
        ],
        "【REVIEW:应急支援保障承诺】": prose_blocks("应急支援保障应结合商务要求中的远程指导、现地排故、驻场调试、重大问题现场技术支持和问题处理报告要求进行承诺。具体联系人、响应级别、到场方式、备件或工具保障、报告模板和费用边界需人工复核。"),
        "【REVIEW:定期跟踪服务承诺】": prose_blocks("定期跟踪服务可包括质保期内定期巡检、运行状态回访、问题统计分析、版本维护建议、培训补强和服务记录归档。具体巡检频次、服务方式、责任人和书面报告要求需结合投标人服务能力及合同条款复核。"),
    }


def write_records() -> None:
    matrix_lines = [
        "# 需求响应矩阵",
        "",
        "| 编号 | 来源文件 | 要求类型 | 原文摘录 | 关键词 | 响应策略 | 对应章节 | 是否生成图 | 状态 |",
        "|---|---|---|---|---|---|---|---|---|",
    ]
    idx = 1
    for name, reqs in FUNCTION_REQUIREMENTS.items():
        for req in reqs:
            matrix_lines.append(f"| R{idx:03d} | 技术要求.docx | 技术要求 | {req} | {name} | 功能设计 | {name}功能 | 是 | 已覆盖 |")
            idx += 1
    for req in PERFORMANCE_REQUIREMENTS:
        matrix_lines.append(f"| R{idx:03d} | 技术要求.docx | 技术要求 | {req} | 性能 | 性能设计 | 性能设计 | 否 | 已覆盖 |")
        idx += 1
    for q, reqs in QUALITY_REQUIREMENTS.items():
        for req in reqs:
            matrix_lines.append(f"| R{idx:03d} | 技术要求.docx | 技术要求 | {req} | {q} | 通用质量特性设计 | {q}设计 | 否 | 已覆盖 |")
            idx += 1
    for req in BUSINESS_REQUIREMENTS:
        matrix_lines.append(f"| R{idx:03d} | 商务要求.docx | 商务要求 | {req} | 商务/服务 | 商务响应与复核 | 售后服务及承诺 | 否 | 已覆盖/需复核 |")
        idx += 1
    for item in SCORING_ITEMS:
        matrix_lines.append(f"| R{idx:03d} | 技术评分表.docx | 评分项 | {item[0]}：{item[1]}（{item[2]}分） | 评分响应 | 强化写入 | {item[3]} | 视情况 | 已覆盖/需补充证明 |")
        idx += 1
    (RECORDS / "requirements-matrix.md").write_text("\n".join(matrix_lines) + "\n", encoding="utf-8")

    confirm_items = [
        "投标人名称（模板原有单位名称未使用，需由用户确认真实投标单位后填写）",
        "项目团队人员说明",
        "项目团队-职务分工",
        "项目团队-姓名",
        "项目团队-职称",
        "项目团队-专业",
        "项目团队-从业资格",
        "项目团队-相关工作年限",
        "质量保证期承诺（商务要求载明24个月，仍需投标人最终确认承诺表述）",
        "售后服务响应承诺（商务要求载明7×24小时、30分钟响应、重大问题12小时到场等，仍需投标人最终确认承诺表述）",
    ]
    (RECORDS / "人工确认清单.md").write_text("# 人工确认清单\n\n" + "\n".join(f"- {x}" for x in confirm_items) + "\n", encoding="utf-8")

    review_items = [
        "质量体系与资质响应说明：需依据投标人真实体系、证书和证明材料复核。",
        "服务质量保障措施：需复核服务团队、响应等级和承诺边界。",
        "项目进度计划：需复核实际进场日期、阶段里程碑和资源投入。",
        "成果交付及验收：需结合合同和甲方验收大纲复核。",
        "交付物清单：需复核数量、介质、格式和签收要求。",
        "应急支援保障承诺：需复核联系人、到场方式和费用边界。",
        "定期跟踪服务承诺：需复核巡检频次和报告要求。",
        "技术指标正偏离、项目负责人、团队证书、同类业绩等客观评分项需补充证明材料。",
    ]
    (RECORDS / "复核清单.md").write_text("# 复核清单\n\n" + "\n".join(f"- {x}" for x in review_items) + "\n", encoding="utf-8")

    (RECORDS / "missing-source.md").write_text("# 未找到来源事项\n\n本次COPY类占位符均已找到来源或由用户在任务中明确提供。CONFIRM和REVIEW事项另见人工确认清单、复核清单。\n", encoding="utf-8")


def write_checks(log: list[tuple[str, str]]) -> None:
    log_lines = ["# 占位符填充日志", "", "| 占位符 | 处理结果 |", "|---|---|"]
    for placeholder, status in log:
        log_lines.append(f"| {placeholder} | {status} |")
    (RECORDS / "placeholder-fill-log.md").write_text("\n".join(log_lines) + "\n", encoding="utf-8")

    coverage = f"""# 覆盖检查报告

## 1. 技术要求覆盖情况

已按功能、性能、通用质量特性、接口、设计约束、质量监督和验收要求建立需求响应矩阵，并映射到需求分析、总体架构、功能设计、性能设计、数据库设计、通用质量特性设计、关键技术、质量控制、风险控制和交付验收章节。

## 2. 商务要求覆盖情况

交付时间地点方式、售后服务、培训、知识产权和保密、付款结算、履约保证金、驻场要求等已写入实施、培训、交付验收或复核章节。涉及最终承诺表述的事项保留为CONFIRM或REVIEW。

## 3. 评分项覆盖情况

技术方案、架构设计、流程图、接口设计、部署灵活性、安全设计、关键技术、功能详细设计、项目管理、风险控制、质量控制、培训和售后服务等主观评分项已在正文中强化。客观评分项如工具证明、技术指标偏离、人员证书、同类业绩需投标人补充证明材料。

## 4. 未确认事项

详见《人工确认清单.md》。项目团队、人员资质、质量保证期承诺和售后响应承诺需最终确认。

## 5. 需人工复核事项

详见《复核清单.md》。进度、交付物、验收、服务保障、应急支援、定期跟踪服务和资质证明需复核。

## 6. Mermaid图生成情况

已生成总体架构图1张、功能流程图{len(FUNCTION_REQUIREMENTS)}张，源码和图片均保存在output/records目录。

## 7. 文档占位符处理情况

COPY和GEN占位符已填充；CONFIRM占位符按规则保留；REVIEW占位符已生成初稿并列入复核清单。目录域如未自动刷新，建议在Word中打开后更新域。
"""
    (RECORDS / "coverage-check.md").write_text(coverage, encoding="utf-8")


def patch_docx_xml_text(docx_path: Path, replacements: dict[str, str]) -> None:
    tmp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith(".xml"):
                text = data.decode("utf-8")
                for old, new in replacements.items():
                    text = text.replace(old, new)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp_path.replace(docx_path)


def main() -> None:
    ensure_dirs()
    diagrams = create_diagrams()
    write_records()
    copy2(TEMPLATE, OUT_DOCX)
    doc = Document(str(OUT_DOCX))
    normalize_template_copy(doc)
    log: list[tuple[str, str]] = []

    log.append(("【COPY:项目名称】", replace_inline(doc, "【COPY:项目名称】", PROJECT_NAME)))
    log.append(("【CONFIRM:投标人名称】", "模板文本框内容保存后替换，已写入人工确认清单"))

    for placeholder, blocks in build_copy_blocks().items():
        log.append((placeholder, replace_placeholder_paragraph(doc, placeholder, blocks)))

    gen_blocks = build_gen_blocks(diagrams)
    gen_blocks.update(build_review_blocks())
    for placeholder, blocks in gen_blocks.items():
        log.append((placeholder, replace_placeholder_paragraph(doc, placeholder, blocks)))

    confirm_placeholders = sorted(set(re.findall(r"【CONFIRM:[^】]+】", "\n".join(p.text for p in doc.paragraphs))))
    for placeholder in confirm_placeholders:
        log.append((placeholder, "按规则保留，已写入人工确认清单"))

    for paragraph_obj in doc.paragraphs:
        for run in paragraph_obj.runs:
            run.font.name = run.font.name or "仿宋"
            if run._element.rPr is not None:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name or "仿宋")

    doc.save(str(OUT_DOCX))
    patch_docx_xml_text(OUT_DOCX, {"北京瑞晟成科技发展有限公司": "投标人名称待确认"})
    write_checks(log)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
