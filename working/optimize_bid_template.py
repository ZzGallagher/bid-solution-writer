from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "templates" / "投标方案模板.docx"
BACKUP = ROOT / "templates" / "投标方案模板.优化前备份.docx"


def paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if paragraph.style:
        new_para.style = paragraph.style
    new_para.add_run(text)
    return new_para


def set_text(paragraph: Paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_paragraph(doc: Document, text: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def replace_exact(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            set_text(paragraph, replacements[text])


def replace_after_heading(doc: Document, heading: str, placeholder: str) -> None:
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() != heading:
            continue
        if index + 1 < len(paragraphs) and not paragraphs[index + 1].text.strip():
            set_text(paragraphs[index + 1], placeholder)
        elif index + 1 < len(paragraphs) and paragraphs[index + 1].text.strip().startswith("【"):
            set_text(paragraphs[index + 1], placeholder)
        else:
            paragraph_after(paragraph, placeholder)
        return


def compress_between(doc: Document, start_heading: str, end_heading: str | None, placeholders: list[str]) -> None:
    paragraphs = doc.paragraphs
    start = next((i for i, p in enumerate(paragraphs) if p.text.strip() == start_heading), None)
    if start is None:
        return
    if end_heading is None:
        end = len(paragraphs)
    else:
        end = next((i for i, p in enumerate(paragraphs[start + 1 :], start + 1) if p.text.strip() == end_heading), len(paragraphs))

    body = paragraphs[start + 1 : end]
    for offset, placeholder in enumerate(placeholders):
        if offset < len(body):
            set_text(body[offset], placeholder)
        else:
            paragraph_after(paragraphs[start + offset], placeholder)

    for paragraph in body[len(placeholders) :]:
        delete_paragraph(paragraph)


def normalize_tables(doc: Document) -> None:
    # Project team rows should stay blank for Word editing, but the first data row
    # carries explicit confirmation markers for automated filling.
    for table in doc.tables:
        if not table.rows:
            continue
        header = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
        if "拟担任" in header and "姓名" in header:
            markers = [
                "【CONFIRM:项目团队-职务分工】",
                "【CONFIRM:项目团队-姓名】",
                "【CONFIRM:项目团队-职称】",
                "【CONFIRM:项目团队-专业】",
                "【CONFIRM:项目团队-从业资格】",
                "【CONFIRM:项目团队-相关工作年限】",
            ]
            for cell, marker in zip(table.rows[1].cells, markers):
                cell.text = marker


def main() -> None:
    if not BACKUP.exists():
        copy2(TEMPLATE, BACKUP)

    doc = Document(str(TEMPLATE))

    replace_exact(
        doc,
        {
            "【GEN:根据技术要求及本项目的项目背景，生成1大段文字描述本文件的编写目的】": "【GEN:编写目的】",
            "【GEN:根据技术要求及本项目的项目背景，生成1大段文字描述项目的建设内容】": "【GEN:建设内容】",
            "【COPY:技术要求-通用质量特性要求/非功能性要求/六性要求】": "【COPY:技术要求-通用质量特性要求】",
            "【GEN:总体架构设计（根据技术要求，采用大段文字的形式描述系统总体架构设计）】": "【GEN:总体架构设计】",
            "【GEN:用大段文字描述架构图】": "【GEN:架构图说明】",
            "【GEN:用大段文字描述前述架构的设计原则】": "【GEN:设计原则】",
            "【GEN:根据技术要求，设计软件系统部署架构，并用大段文字描述】": "【GEN:部署架构设计】",
            "【性能设计章节需要一一对应技术要求中的性能需求逐条回答】": "【GEN:性能设计章节】",
        },
    )

    compress_between(
        doc,
        "功能设计",
        "性能设计",
        [
            "【GEN:功能设计总述】",
            "【GEN:功能设计章节】",
        ],
    )

    replace_after_heading(doc, "性能设计", "【GEN:性能设计章节】")
    replace_after_heading(doc, "数据库设计", "【GEN:数据库设计总述】")
    replace_after_heading(doc, "数据库架构", "【GEN:数据库架构设计】")
    replace_after_heading(doc, "核心业务设计", "【GEN:核心业务数据设计】")
    replace_after_heading(doc, "数据库表设计", "【GEN:数据库表设计】")
    replace_after_heading(doc, "通用质量特性设计", "【GEN:通用质量特性设计总述】")
    replace_after_heading(doc, "可靠性设计", "【GEN:可靠性设计】")
    replace_after_heading(doc, "维修性设计", "【GEN:维修性设计】")
    replace_after_heading(doc, "保障性设计", "【GEN:保障性设计】")
    replace_after_heading(doc, "测试性设计", "【GEN:测试性设计】")
    replace_after_heading(doc, "安全性设计", "【GEN:安全性设计】")
    replace_after_heading(doc, "环境适应性设计", "【GEN:环境适应性设计】")
    replace_after_heading(doc, "关键技术", "【GEN:关键技术】")
    replace_after_heading(doc, "质量控制", "【GEN:质量控制总述】")
    replace_after_heading(doc, "风险评估与控制", "【GEN:风险评估与控制】")

    compress_between(
        doc,
        "质量保证措施",
        "项目进度计划",
        [
            "【GEN:质量保证措施】",
            "【REVIEW:质量体系与资质响应说明】",
            "【REVIEW:服务质量保障措施】",
        ],
    )
    compress_between(doc, "项目进度计划", "项目人员组成", ["【REVIEW:项目进度计划】"])
    compress_between(doc, "项目人员组成", "成果交付及验收", ["【CONFIRM:项目团队人员说明】"])
    compress_between(
        doc,
        "成果交付及验收",
        "售后服务及承诺",
        [
            "【REVIEW:成果交付及验收】",
            "【REVIEW:交付物清单】",
        ],
    )
    compress_between(
        doc,
        "售后服务及承诺",
        None,
        [
            "【GEN:培训方案】",
            "【CONFIRM:质量保证期承诺】",
            "【CONFIRM:售后服务响应承诺】",
            "【REVIEW:应急支援保障承诺】",
            "【REVIEW:定期跟踪服务承诺】",
        ],
    )

    normalize_tables(doc)
    doc.save(str(TEMPLATE))


if __name__ == "__main__":
    main()
